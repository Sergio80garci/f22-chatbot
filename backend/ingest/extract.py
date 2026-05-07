import os
from pathlib import Path
from typing import Optional

import pypdf
import pytesseract
from PIL import Image
from docx import Document
from openpyxl import load_workbook

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".png", ".jpg", ".jpeg", ".txt"}
OCR_MIN_CHARS_PER_PAGE = 100


def _extract_pdf(path: Path) -> dict:
    pages_text = []
    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)
        total_pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if len(text.strip()) < OCR_MIN_CHARS_PER_PAGE:
                # Fallback to OCR for scanned pages
                try:
                    img = page.images[0].image if page.images else None
                    if img:
                        text = pytesseract.image_to_string(img, lang="spa")
                except Exception:
                    pass
            pages_text.append(text)
    return {
        "filename": path.name,
        "content": "\n".join(pages_text),
        "metadata": {
            "source": str(path),
            "file_type": "pdf",
            "total_pages": total_pages,
        },
    }


def _extract_docx(path: Path) -> dict:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return {
        "filename": path.name,
        "content": "\n".join(parts),
        "metadata": {"source": str(path), "file_type": "docx"},
    }


def _extract_xlsx(path: Path) -> dict:
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Hoja: {sheet.title}]")
        headers = None
        for row in sheet.iter_rows(values_only=True):
            if headers is None:
                headers = [str(c) if c is not None else "" for c in row]
                parts.append("\t".join(headers))
            else:
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    parts.append(line)
    return {
        "filename": path.name,
        "content": "\n".join(parts),
        "metadata": {"source": str(path), "file_type": "xlsx"},
    }


def _extract_image(path: Path) -> dict:
    img = Image.open(path)
    text = pytesseract.image_to_string(img, lang="spa")
    return {
        "filename": path.name,
        "content": text,
        "metadata": {"source": str(path), "file_type": path.suffix.lstrip(".")},
    }


def _extract_txt(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    return {
        "filename": path.name,
        "content": text,
        "metadata": {"source": str(path), "file_type": "txt"},
    }


def extract_file(file_path: str) -> Optional[dict]:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        print(f"  [SKIP] Formato no soportado: {path.name}")
        return None
    print(f"  Extrayendo: {path.name}")
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".xlsx":
            return _extract_xlsx(path)
        elif ext in {".png", ".jpg", ".jpeg"}:
            return _extract_image(path)
        elif ext == ".txt":
            return _extract_txt(path)
    except Exception as e:
        print(f"  [ERROR] {path.name}: {e}")
        return None
